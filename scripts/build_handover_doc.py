from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[1]
SOURCE = PROJECT_ROOT / "docs" / "Vision-Platform-System-Handover.md"
OUTPUT = PROJECT_ROOT / "docs" / "Vision-Platform-System-Handover.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "203748"
MUTED = "5C6773"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
BORDER = "CCD5DF"
CODE_FILL = "F6F8FA"
BODY_FONT = "Calibri"
CJK_FONT = "Microsoft YaHei"
CONTENT_WIDTH_DXA = 9360


def set_run_font(
    run,
    *,
    name: str = BODY_FONT,
    size: float | None = None,
    color: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), CJK_FONT)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = document.styles[style_name]
        style.font.name = BODY_FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = document.styles[style_name]
        style.font.name = BODY_FONT
        style.font.size = Pt(11)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            cell.width = Inches(widths[index] / 1440)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            margins = tc_pr.find(qn("w:tcMar"))
            if margins is None:
                margins = OxmlElement("w:tcMar")
                tc_pr.append(margins)
            for edge, value in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
                node = margins.find(qn(f"w:{edge}"))
                if node is None:
                    node = OxmlElement(f"w:{edge}")
                    margins.append(node)
                node.set(qn("w:w"), str(value))
                node.set(qn("w:type"), "dxa")


def add_field(paragraph, instruction: str) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()
    run._r.extend((begin, instr, separate, text, end))


def configure_page(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run("VISION PLATFORM  |  SYSTEM HANDOVER")
    set_run_font(run, size=8.5, color=MUTED, bold=True)

    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("内部交接文档  ·  第 ")
    set_run_font(run, size=8.5, color=MUTED)
    add_field(paragraph, "PAGE")
    run = paragraph.add_run(" 页")
    set_run_font(run, size=8.5, color=MUTED)


def add_cover(document: Document) -> None:
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(110)

    kicker = document.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = kicker.add_run("工业视觉智能平台  ·  系统交接")
    set_run_font(run, size=11, color=BLUE, bold=True)
    kicker.paragraph_format.space_after = Pt(18)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Vision Platform")
    set_run_font(run, size=30, color=NAVY, bold=True)
    title.paragraph_format.space_after = Pt(8)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("工业视觉智能平台系统交接手册")
    set_run_font(run, size=16, color=DARK_BLUE, bold=True)
    subtitle.paragraph_format.space_after = Pt(26)

    description = document.add_paragraph()
    description.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = description.add_run("功能 · 实现 · 架构 · 接口 · 部署 · 运维 · 风险")
    set_run_font(run, size=11, color=MUTED)
    description.paragraph_format.space_after = Pt(90)

    metadata = document.add_table(rows=4, cols=2)
    metadata.style = "Table Grid"
    rows = (
        ("文档版本", "V1.0"),
        ("对应分支", "develop"),
        ("编制日期", "2026-08-16"),
        ("适用对象", "开发、算法、实施、运维与项目负责人"),
    )
    for row, values in zip(metadata.rows, rows):
        for cell, value in zip(row.cells, values):
            cell.text = value
        set_cell_shading(row.cells[0], LIGHT_BLUE)
        for run in row.cells[0].paragraphs[0].runs:
            set_run_font(run, size=10, bold=True, color=DARK_BLUE)
        for run in row.cells[1].paragraphs[0].runs:
            set_run_font(run, size=10)
    set_table_geometry(metadata, [2700, 6660])
    document.add_page_break()


def add_static_toc(document: Document, headings: list[str]) -> None:
    paragraph = document.add_paragraph("目录", style="Heading 1")
    paragraph.paragraph_format.space_after = Pt(12)
    for heading in headings:
        item = document.add_paragraph()
        item.paragraph_format.left_indent = Inches(0.15)
        item.paragraph_format.space_after = Pt(2)
        run = item.add_run(heading)
        set_run_font(run, size=10.2, color=NAVY)
    document.add_page_break()


INLINE_PATTERN = re.compile(r"(\*\*.+?\*\*|`.+?`)")


def add_inline(paragraph, text: str) -> None:
    position = 0
    for match in INLINE_PATTERN.finditer(text):
        if match.start() > position:
            run = paragraph.add_run(text[position : match.start()])
            set_run_font(run)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, bold=True)
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, name="Consolas", size=9.5, color=DARK_BLUE)
        position = match.end()
    if position < len(text):
        run = paragraph.add_run(text[position:])
        set_run_font(run)


def add_code_block(document: Document, lines: list[str]) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.18)
    paragraph.paragraph_format.right_indent = Inches(0.18)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.05
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), CODE_FILL)
    p_pr.append(shading)
    run = paragraph.add_run("\n".join(lines))
    set_run_font(run, name="Consolas", size=8.8, color="263238")


def add_table(document: Document, rows: list[list[str]]) -> None:
    if len(rows) < 2:
        return
    column_count = len(rows[0])
    table = document.add_table(rows=len(rows), cols=column_count)
    table.style = "Table Grid"
    for row_index, values in enumerate(rows):
        for column_index, value in enumerate(values):
            cell = table.cell(row_index, column_index)
            cell.text = ""
            paragraph = cell.paragraphs[0]
            add_inline(paragraph, value.strip())
            for run in paragraph.runs:
                set_run_font(
                    run,
                    size=9.2,
                    bold=row_index == 0,
                    color=DARK_BLUE if row_index == 0 else None,
                )
            if row_index == 0:
                set_cell_shading(cell, LIGHT_BLUE)
            elif row_index % 2 == 0:
                set_cell_shading(cell, "FAFBFC")

    if column_count == 2:
        widths = [2700, 6660]
    elif column_count == 3:
        widths = [2000, 2780, 4580]
    elif column_count == 4:
        widths = [1250, 1900, 2650, 3560]
    else:
        base = CONTENT_WIDTH_DXA // column_count
        widths = [base] * column_count
        widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    set_table_geometry(table, widths)
    document.add_paragraph().paragraph_format.space_after = Pt(1)


def parse_markdown(document: Document, text: str) -> None:
    lines = text.splitlines()
    headings = [line[3:].strip() for line in lines if line.startswith("## ")]
    add_static_toc(document, headings)

    index = 0
    paragraph_buffer: list[str] = []
    in_code = False
    code_lines: list[str] = []

    def flush_paragraph() -> None:
        if not paragraph_buffer:
            return
        paragraph = document.add_paragraph()
        add_inline(paragraph, " ".join(item.strip() for item in paragraph_buffer))
        paragraph_buffer.clear()

    while index < len(lines):
        line = lines[index]
        stripped = line.strip()

        if stripped.startswith("```"):
            flush_paragraph()
            if in_code:
                add_code_block(document, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            index += 1
            continue
        if in_code:
            code_lines.append(line)
            index += 1
            continue

        if not stripped:
            flush_paragraph()
            index += 1
            continue

        if stripped == "---":
            flush_paragraph()
            index += 1
            continue

        if line.startswith("# ") or line.startswith("> 文档版本") or line.startswith("> 对应代码") or line.startswith("> 编制日期") or line.startswith("> 适用对象"):
            index += 1
            continue

        if line.startswith("## "):
            flush_paragraph()
            document.add_paragraph(line[3:].strip(), style="Heading 1")
            index += 1
            continue
        if line.startswith("### "):
            flush_paragraph()
            document.add_paragraph(line[4:].strip(), style="Heading 2")
            index += 1
            continue
        if line.startswith("#### "):
            flush_paragraph()
            document.add_paragraph(line[5:].strip(), style="Heading 3")
            index += 1
            continue

        if stripped.startswith("|") and index + 1 < len(lines) and re.match(r"^\|?\s*:?-+", lines[index + 1].strip("| ")):
            flush_paragraph()
            table_rows: list[list[str]] = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                values = [value.strip() for value in lines[index].strip().strip("|").split("|")]
                if not all(re.fullmatch(r":?-+:?", value.replace(" ", "")) for value in values):
                    table_rows.append(values)
                index += 1
            add_table(document, table_rows)
            continue

        if re.match(r"^- \[ \] ", stripped):
            flush_paragraph()
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.18)
            paragraph.paragraph_format.space_after = Pt(4)
            run = paragraph.add_run("☐ ")
            set_run_font(run, size=11, color=BLUE)
            add_inline(paragraph, stripped[6:])
            index += 1
            continue

        if stripped.startswith("- "):
            flush_paragraph()
            paragraph = document.add_paragraph(style="List Bullet")
            add_inline(paragraph, stripped[2:])
            index += 1
            continue

        if re.match(r"^\d+\. ", stripped):
            flush_paragraph()
            paragraph = document.add_paragraph(style="List Number")
            add_inline(paragraph, re.sub(r"^\d+\. ", "", stripped))
            index += 1
            continue

        if stripped.startswith("> "):
            flush_paragraph()
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.25)
            paragraph.paragraph_format.right_indent = Inches(0.25)
            paragraph.paragraph_format.space_before = Pt(4)
            paragraph.paragraph_format.space_after = Pt(6)
            p_pr = paragraph._p.get_or_add_pPr()
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), LIGHT_GRAY)
            p_pr.append(shading)
            add_inline(paragraph, stripped[2:])
            index += 1
            continue

        paragraph_buffer.append(line)
        index += 1

    flush_paragraph()


def build() -> Path:
    markdown = SOURCE.read_text(encoding="utf-8")
    document = Document()
    configure_styles(document)
    configure_page(document)
    add_cover(document)
    parse_markdown(document, markdown)

    for paragraph in document.paragraphs:
        if paragraph.style.name.startswith("Heading"):
            paragraph.paragraph_format.keep_with_next = True
            paragraph.paragraph_format.keep_together = True

    properties = document.core_properties
    properties.title = "Vision Platform 工业视觉智能平台系统交接手册"
    properties.subject = "系统功能、实现、架构、接口、部署和运维交接"
    properties.author = "Vision Platform Project"
    properties.keywords = "工业视觉, FastAPI, DINOv2, Qwen3-VL, Grounding DINO, SAM2, PaddleOCR"

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build())
